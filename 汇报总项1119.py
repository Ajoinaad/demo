#!/usr/bin/env python
# coding: utf-8
import shutil
import os
import time
import datetime
import warnings
from openpyxl import Workbook
from openpyxl.styles import Border, Side, PatternFill, Font, Alignment
import pandas as pd
from openpyxl import load_workbook
import numpy as np
import openpyxl
from docx import Document
from docx.oxml.ns import qn
from openpyxl.utils import get_column_letter

now_time = datetime.datetime.now()
warnings.filterwarnings('ignore')
name = ' '  # 所有病例数据处理
title = '请输入抬头'  # 默认抬头


# 先计算密接类型
class Original:
    report = pd.read_excel('批量汇报.xlsx', sheet_name='汇报专项')  # 读取汇报模板
    doc = Document()

    # 总数
    def __init__(self, name='病例名称'):
        self.name = name

    # 1.文件名
    def open_file(self):
        # print('-' * 80)
        files = r"input"  # 打开文件夹
        lists = os.listdir(files)  # 列出目录的下所有文件和文件夹保存到lists
        lists.sort(key=lambda fn: os.path.getmtime(files + "\\" + fn))  # 按时间排序
        file_new = os.path.join(files, lists[-1])
        return file_new

    # 2.数据处理
    def dealwith(self):
        print('\n\n\n')
        print('数据正在处理中...')
        # 0.删除旧文件,创建新文件
        try:
            shutil.rmtree(rf'D:\工作源文件')
        except:
            print('数据处理出现异常,请关闭所有文件夹/文件,并检查D:\盘是否有D:\工作源文件,有则删除,没有请重新运行')

        os.makedirs(rf'D:\工作源文件')
        # 1.读取最新需要处理文件
        df_deal = pd.read_excel(f'./{Original().open_file()}')

        # 1.1重复值处理(目前还不确定是否要弄)
        # df_deal = df.drop_duplicates(subset=['有效证件号','是否排除密接/次密','密接类型','转归'], keep='first', inplace=True)
        # 2.数据空值处理
        df_deal['审核时间'].fillna('空白', inplace=True)
        df_deal['镇（街道）'].fillna('空白', inplace=True)
        df_deal['转出目的省(直辖市)'].fillna(df_deal['目前所处位置'], inplace=True)
        df_deal['转出目的省(直辖市)'].fillna(df_deal['现住址'], inplace=True)
        df_deal['转出目的省(直辖市)'].fillna('不明', inplace=True)
        df_deal['转归'].fillna('空白', inplace=True)
        df_deal['医学观察场所名称'].fillna('空白', inplace=True)
        # 街道错误处理
        road_List = ['同德', '松洲', '黄石', '石井', '鹤龙', '龙归', '金沙', '嘉禾', '云城', '三元里', '京溪', '同和', '人和', '均禾', '大源', '太和',
                     '白云湖', '景泰', '棠景', '永平', '江高', '石门', '新市', '钟落潭']
        name_List = name.split()
        # name_List = name.split('、')
        name_type = df_deal['关联病例'].str.contains('|'.join(name_List))
        must = (df_deal['是否排除密接/次密'] != '是') & (df_deal['是否追踪到'] != '转出外省') & (df_deal['地市'] == '广州市') & (
                df_deal['区县'] == '白云区') & name_type
        worng_must = must & (df_deal['镇（街道）'].str.contains('|'.join(road_List)) == False)
        df_deal['镇（街道）'].mask(worng_must, '待甄别街道', inplace=True)
        # 2.1转换时间控制处理:待转运处理
        df_deal['一码通入住日期待转运'] = df_deal['一码通入住日期'].fillna('空白')
        df_deal['医学观察方式'].mask(((df_deal['一码通入住日期待转运'] == '空白') & (df_deal['居家隔离原因'].isnull())), '待转运',
                               inplace=True)
        # 2.1待转运处理（1018更新）
        df_deal['医学观察方式'].mask(df_deal['医学观察场所名称'].str.contains('待转运'), '待转运', inplace=True)

        # 2.2日期处理
        df_deal['一码通入住日期'].fillna('1970-12-30 23:59', inplace=True)
        df_deal['最后接触日期'].fillna('1970-12-31 23:59', inplace=True)
        df_deal['一码通入住日期'] = pd.to_datetime(df_deal['一码通入住日期']).dt.floor('d')
        df_deal['最后接触日期'] = pd.to_datetime(df_deal['最后接触日期']).dt.floor('d')
        # 2.3数据空值预处理
        df_deal['审核时间'].fillna('空白', inplace=True)
        df_deal['镇（街道）'].fillna('空白', inplace=True)
        df_deal['是否核心密接'].fillna('否', inplace=True)
        # 2.4集中处理
        # 最后接触日期<一码通入住日期就改为集中
        df_deal['医学观察方式'].mask(((df_deal['最后接触日期']) < (df_deal['一码通入住日期'])), '集中', inplace=True)
        df_deal['医学观察方式'].mask(((df_deal['医学观察场所名称'].str.contains('酒店')) & (df_deal['目前所处位置'] != df_deal['医学观察场所名称'])),
                               '集中', inplace=True)
        # 375个场所

        # 3.确诊处理
        df_deal['转归'] = df_deal['转归'].str.replace(pat='.*确诊.*', repl='转为确诊', regex=True)
        df_deal['转归'] = df_deal['转归'].str.replace(pat='.*阳性.*', repl='转为确诊', regex=True)

        # 4.10月5日更新：次密：医学观察方式为“待转运”皆归为“居家”（并把居家隔离原因里的待转运删除）
        df_deal['医学观察方式'].mask((df_deal['医学观察方式'].str.contains('待转运')) & (df_deal['密接类型'] == '密接的密接'), '居家',
                               inplace=True)
        df_deal['居家隔离原因'].mask((df_deal['医学观察方式'].str.contains('待转运')) & (df_deal['密接类型'] == '密接的密接'), '空白',
                               inplace=True)

        df_deal = df_deal.get(
            ['地市', '区县', '镇（街道）', 'ID', '姓名', '国籍', '性别', '年龄', '有效证件号', '联系方式', '目前所处位置', '现住址', '职业', '工作单位',
             '密接/次密发现途径', '是否核心密接', '是否排除密接/次密', '关联病例', '关联密接', '密接类型', '与患者关系', '接触地点', '最后接触日期', '应解除观察日期', '关联重点场所',
             '转归', '备注', '是否追踪到', '审核时间', '录入时间', '医学观察方式', '创建单位', '转出目的省(直辖市)', '居家隔离原因'])

        # 加快速度专用！这样排除密接/次密身份!!!!!(这一句大疫情过后可以删除)
        df_deal = df_deal[df_deal['是否排除密接/次密'] != '是']
        # 新模板删除次密部分（11.12日以后）
        df_deal = df_deal[df_deal['密接类型'] != '密接的密接']

        #
        df_deal.to_csv(f'D:/工作源文件/{Original().open_file()[5:-5]}.csv', index=False)  # 保存最新版数据
        # df_deal.to_excel(f'D:/工作源文件/{Original().open_file()[5:]}', index=False)  # 保存最新版数据
        print('数据处理完成!')

    # 3.文字报数（11.12大更新）
    @classmethod
    def report_pronusual(self):
        print('\n\n\n')
        # 打印标题
        y = f'未审核\n★内部材料 注意保密 外传必究★\n【白云区密接录入及管控情况通报，{now_time.month}月{now_time.day}日{now_time.hour + 1}时00分】（下库时间：{Original().open_file()[12:-5]} 以下数据系动态变化）'
        print(y)
        z = f'【{title}涉我区密接管控情况】（数据来源于省流调系统）'
        print(z)
        print('')
        # word打印标题
        Original().doc.add_heading(f'【{title}】', level=1)
        text = y + '\n' + z
        Original().doc.add_paragraph(text)
        # 读取处理文件
        df = pd.read_csv(f'D:/工作源文件/{Original().open_file()[5:-5]}.csv')
        # 名字类别
        name_List = name.split()
        name_type = df['关联病例'].str.contains('|'.join(name_List))
        # 上一版日期和上一版数据
        path2 = r"./last"
        df_last = os.listdir(path2)
        last = os.path.join(path2, df_last[-1])
        df_lastexcel = pd.read_excel(f'{last}')
        # 均生成核心密接
        type = ['是', '']
        sick_type = '核心密接'
        # 是否和上版人数对比
        if lastnameif == '是':
            last_name = name_last.split()
            last_name_type = df_lastexcel['关联病例'].str.contains('|'.join(last_name))
        else:
            last_name = name.split()
            last_name_type = df_lastexcel['关联病例'].str.contains('|'.join(last_name))
        for i in range(0, len(type)):
            # 插入:统计第四点:先统计第四点
            type_type = type[i]

            # 上一版我区应管核密/密接人数
            if sick_type == '密接':
                mi = (df_lastexcel['是否追踪到'] != '转出外省') & (df_lastexcel['地市'] == '广州市') & (df_lastexcel['区县'] == '白云区') & (df_lastexcel['是否排除密接/次密'] != '是') & last_name_type & (df_lastexcel['密接类型'] == '密切接触者')
                count_last = len(df_lastexcel[mi])
            else:
                mi = (df_lastexcel['是否追踪到'] != '转出外省') & (df_lastexcel['地市'] == '广州市') & (df_lastexcel['区县'] == '白云区') & (df_lastexcel['是否排除密接/次密'] != '是') & (df_lastexcel['是否核心密接'].str.contains('|'.join(type[i]))) & last_name_type & (df_lastexcel['密接类型'] == '密切接触者')
                count_last = len(df_lastexcel[mi])
            # 第一点
            # 累计甄别
            count = (df['是否排除密接/次密'] != '是') & (df['是否核心密接'].str.contains('|'.join(type[i]))) & name_type & (df['密接类型'] == '密切接触者')
            count_str = f'1.累计甄别{sick_type}{len(df[count])}人：'
            # 我区主动甄别
            active = count & (df['创建单位'] == '白云区疾病预防控制中心')
            active_count = len(df[active])
            if active_count == 0:
                active_str = f'均为区外推送；'
            else:
                active_str = f'其中我区主动甄别{active_count}人【'

            # 推送外省管控人数
            province = active & (df['是否追踪到'] == '转出外省')
            province_count = len(df[province])
            # 装载细分
            partdict = {}
            partlist = []
            partstr = ''
            for i in df[province]['转出目的省(直辖市)']:
                if (i[0:2] == '内蒙') or (i[0:2] == '黑龙'):
                    partlist.append(i[0:3])
                else:
                    partlist.append(i[0:2])
            for i in partlist:
                partdict[i] = partdict.get(i, 0) + 1
            partsort = sorted(partdict.items(), key=lambda x: x[1], reverse=True)
            for i in range(0, len(partsort)):
                if partsort[i][1] == 1:
                    partstr += f'{partsort[i][0]}、'
                else:
                    partstr += f'{partsort[i][0]}{partsort[i][1]}人，'

            if province_count == 0:
                province_str = ''
            elif province_count == 1:
                province_str = f'推送外省{province_count}人({partstr[:-1]}1人)，'
            elif partstr[-2] == '人':
                province_str = f'推送外省{province_count}人({partstr[:-1]})，'
            elif partsort[-1][1] != partsort[-2][1]:
                province_str = f'推送外省{province_count}人({partstr[:-1]}1人)，'
            else:
                province_str = f'推送外省{province_count}人({partstr[:-1]}各1人)，'

            # 推送外市管控人数
            city = active & (df['是否追踪到'] != '转出外省') & (df['地市'] != '广州市')
            city_count = len(df[city])
            # 装载细分
            partdict = {}
            partlist = []
            partstr = ''
            for i in df[city]['地市']:
                partlist.append(i)
            for i in partlist:
                partdict[i] = partdict.get(i, 0) + 1
            partsort = sorted(partdict.items(), key=lambda x: x[1], reverse=True)
            for i in range(0, len(partsort)):
                if partsort[i][1] == 1:
                    partstr += f'{partsort[i][0]}、'
                else:
                    partstr += f'{partsort[i][0]}{partsort[i][1]}人，'

            if city_count == 0:
                city_str = ''
            elif city_count == 1:
                city_str = f'推送外市{city_count}人({partstr[:-1]}1人)，'
            elif partstr[-2] == '人':
                city_str = f'推送外市{city_count}人({partstr[:-1]})，'
            elif partsort[-1][1] != partsort[-2][1]:
                city_str = f'推送外市{city_count}人({partstr[:-1]}1人)，'
            else:
                city_str = f'推送外市{city_count}人({partstr[:-1]}各1人)，'
            # 推送外区管控人数
            area = active & (df['是否追踪到'] != '转出外省') & (df['地市'] == '广州市') & (df['区县'] != '白云区')
            area_count = len(df[area])
            # 装载细分
            partdict = {}
            partlist = []
            partstr = ''
            for i in df[area]['区县']:
                partlist.append(i)
            for i in partlist:
                partdict[i] = partdict.get(i, 0) + 1
            partsort = sorted(partdict.items(), key=lambda x: x[1], reverse=True)
            for i in range(0, len(partsort)):
                if partsort[i][1] == 1:
                    partstr += f'{partsort[i][0]}、'
                else:
                    partstr += f'{partsort[i][0]}{partsort[i][1]}人，'

            if area_count == 0:
                area_str = ''
            elif area_count == 1:
                area_str = f'推送外区{area_count}人({partstr[:-1]}1人)，'
            elif partstr[-2] == '人':
                area_str = f'推送外区{area_count}人({partstr[:-1]})，'
            elif partsort[-1][1] != partsort[-2][1]:
                area_str = f'推送外区{area_count}人({partstr[:-1]}1人)，'
            else:
                area_str = f'推送外区{area_count}人({partstr[:-1]}各1人)，'

            local = active & (df['是否追踪到'] != '转出外省') & (df['地市'] == '广州市') & (df['区县'] == '白云区')
            local_count = len(df[local])
            if local_count == 0:
                local_str = ''
            else:
                local_str = f'归属白云区{local_count}人】；'

            # 区外推送直接计算
            pushme_count = len(df[count]) - len(df[active])
            if pushme_count == 0:
                pushme_str = ''
            else:
                pushme_str = f'区外推送我区{pushme_count}人；'
            # 外区推送我区,我区推送外省(求Y的值) :区外推送X人（推送外省'Y'人）
            out = count & (df['创建单位'] != '白云区疾病预防控制中心') & (df["是否追踪到"] == "转出外省")
            out_count = len(df[out])
            if out_count != 0:
                out_str = f'我区转推送外省{out_count}人；'
                df[out].to_excel(rf'./output/次要名单/{sick_type}我区转推送外省名单.xlsx', index=False)
            else:
                out_str = ''
            # 第二点
            # 涉及我区应管
            must = count & (df['是否追踪到'] != '转出外省') & (df['地市'] == '广州市') & (df['区县'] == '白云区')
            must_count = len(df[must])
            if must_count == 0:
                must_str = f'暂无{sick_type}涉及我区；'
                simple_must_str = f'暂无{sick_type}涉及我区；'
            else:
                must_str = f'2.涉及我区{sick_type}{must_count}人：'
                simple_must_str = f'涉及我区{sick_type}{must_count}人：'
            # 已落地
            workable = (must & (df['审核时间'] != '空白')) + (must & (df['转归'] == '转为确诊'))
            workable_count = len(df[workable])
            if workable_count == 0:
                workable_str = ''
            else:
                workable_str = f'已落地{workable_count}人；'
            # 目前在管
            atpresent = workable & (df['转归'].str.contains('继续'))
            atpresent_count = len(df[atpresent])
            if atpresent_count == 0:
                atpresent_str = ''
            else:
                atpresent_str = f'目前在管{atpresent_count}人'
            # 集中
            focus = atpresent & (df['医学观察方式'].str.contains('集中'))
            focus_count = len(df[focus])
            if focus_count == 0:
                focus_str = ''
            else:
                focus_str = f'集中隔离{focus_count}人，'
            # 待转运
            wait = atpresent & (df['医学观察方式'].str.contains('待转运'))
            wait_count = len(df[wait])
            if wait_count == 0:
                simple_wait_str = ''
            else:
                simple_wait_str = f'待转运{wait_count}人，'
            partdict = {}
            partlist = []
            partstr = ''
            for i in df[wait]['镇（街道）']:
                partlist.append(i)
            for i in partlist:
                partdict[i] = partdict.get(i, 0) + 1
            partsort = sorted(partdict.items(), key=lambda x: x[1], reverse=True)
            for i in range(0, len(partsort)):
                if partsort[i][1] == 1:
                    partstr += f'{partsort[i][0]}、'
                else:
                    partstr += f'{partsort[i][0]}{partsort[i][1]}人，'
            if wait_count == 0:
                wait_str = ''
            elif wait_count == 1:
                wait_str = f'待转运{wait_count}人({partstr[:-1]}1人)，'
            elif partstr[-2] == '人':
                wait_str = f'待转运{wait_count}人({partstr[:-1]})，'
            elif partsort[-1][1] != partsort[-2][1]:
                wait_str = f'待转运{wait_count}人({partstr[:-1]}1人)，'
            else:
                wait_str = f'待转运{wait_count}人({partstr[:-1]}各1人)，'

            # 居家
            home = atpresent & (df['医学观察方式'].str.contains('居家'))
            home_count = len(df[home])
            if home_count == 0:
                home_str = ''
            else:
                home_str = f'居家隔离{home_count}人'
            # 医院隔离
            hospital = atpresent & (df['医学观察方式'].str.contains('医院'))
            hospital_count = len(df[hospital])
            if hospital_count == 0:
                hospital_str = ''
            else:
                hospital_str = f'，医院隔离{hospital_count}人，'
            # 解除隔离
            remove = must & (df['转归'].str.contains('解除')) & (df['审核时间'] != '空白')
            remove_count = len(df[remove])
            if remove_count == 0:
                remove_str = ''
            else:
                remove_str = f'解除隔离{remove_count}人；'
            # 正在落地管控
            track = must & (df['审核时间'] == '空白') & (df['转归'] != '转为确诊')
            track_count = len(df[track])
            # 简版
            simple_track_str = f'正在落地管控{track_count}人。'
            if track_count == 0:
                simple_track_str = ''
            else:
                simple_track_str = f'正在落地管控{track_count}人。'
            partdict = {}
            partlist = []
            partstr = ''
            for i in df[track]['镇（街道）']:
                if (i[0:2] == '白云') or (i[0:2] == '三元') or (i[0:2] == '待甄') or (i[0:2] == '太和') or (i[0:2] == '江高') or (i[0:2] == '人和'):
                    partlist.append(i[0:3])
                else:
                    partlist.append(i[0:2])
            for i in partlist:
                partdict[i] = partdict.get(i, 0) + 1
            partsort = sorted(partdict.items(), key=lambda x: x[1], reverse=True)
            for i in range(0, len(partsort)):
                if partsort[i][1] == 1:
                    partstr += f'{partsort[i][0]}街道、'
                else:
                    partstr += f'{partsort[i][0]}街道{partsort[i][1]}人，'
            partstr = partstr.replace('江高镇街道', '江高镇')
            partstr = partstr.replace('人和镇街道', '人和镇')
            partstr = partstr.replace('太和镇街道', '太和镇')
            partstr = partstr.replace('钟落街道', '钟落潭镇')
            if track_count == 0:
                track_str = ''
            elif track_count == 1:
                track_str = f'正在落地管控{track_count}人({partstr[:-1]}1人)；'
            elif partstr[-2] == '人':
                track_str = f'正在落地管控{track_count}人({partstr[:-1]})；'
            elif partsort[-1][1] != partsort[-2][1]:
                track_str = f'正在落地管控{track_count}人({partstr[:-1]}1人)；'
            else:
                track_str = f'正在落地管控{track_count}人({partstr[:-1]}各1人)；'

            # 转为确诊
            sun = must & (df['转归'] == '转为确诊')
            sun_count = len(df[sun])
            if sun_count == 0:
                sun_str = ''
            else:
                sun_str = f'转为确诊{sun_count}人；'

            # 第三点，累计已推送区外管控人数（直接计算）
            allpush_count = province_count + city_count + area_count + out_count
            allpush_str = f'3.累计已推送区外管控{allpush_count}人，已完成双握手；'

            # 第四点
            add_count = must_count - count_last
            # 简版
            # 文字质控（1115日更新）
            title_str = f'【{sick_type}管控情况】'
            first_str = count_str + active_str + area_str + city_str + province_str + local_str + pushme_str + out_str
            second_str = must_str + workable_str + atpresent_str + '【' + focus_str + wait_str + home_str + hospital_str + '】；' + remove_str + sun_str + track_str
            second_str = second_str.replace('，】；', '】；')
            third_str = allpush_str
            four_str = f'4.新增{sick_type}：{last[7:-5]}-{Original().open_file()[6:-5]}新增{add_count}人。'
            if must_count != workable_count + track_count:
                wrong_str = f'{sick_type}我区应管与已落地+正在落地管控人数不闭环,请检查原因~'
                print(wrong_str)
            if atpresent_count != focus_count + wait_count + home_count + hospital_count:
                wrong_str = f'{sick_type}目前在管人数与(集中隔离+待转运+居家+医院)人数不闭环,请检查原因~'
                print(wrong_str)
            if workable_count != atpresent_count + remove_count + sun_count:
                wrong_str = f'{sick_type}已落地人数与(目前在管+解除隔离+转为确诊)人数不闭环,请检查原因~'
                print(wrong_str)

            if add_count == 0:
                four_str = four_str.replace('新增0人', '无新增')
            simple_second_str = simple_must_str + workable_str + atpresent_str + '【' + focus_str + simple_wait_str + home_str + hospital_str + '】；' + remove_str + sun_str + simple_track_str
            simple_second_str = simple_second_str.replace('，】；','】；')

            # 打印详细版

            print(title_str)
            print(first_str)
            print(second_str)
            print(third_str)
            Original().doc.add_paragraph(title_str)
            Original().doc.add_paragraph(first_str)
            Original().doc.add_paragraph(second_str)
            Original().doc.add_paragraph(third_str)
            # 第四点，已完善核减情况
            if add_count < 0:
                # 如果有核减
                four_str = f'{four_str}<<==核减和新增必须等于这个数'
                Original().doc.add_paragraph(four_str)
                print(four_str, '<<==核减和新增必须等于这个数')
                # 找出差异的ID
                reduce = pd.concat([df_lastexcel[mi]['ID'], df[must]['ID']]).drop_duplicates(keep=False)
                # 保存差异的ID.excel
                reduce.to_excel(f'./output/重要名单/核减专属名单/4.{sick_type}差异ID.xlsx', index=False)
                # 读取这个IDExcel
                reduce = pd.read_excel(f'./output/重要名单/核减专属名单/4.{sick_type}差异ID.xlsx')
                # 读取原始input的excel
                up_to_date = pd.read_excel(f'./{Original().open_file()}')
                # 两表的差异
                df_check = pd.merge(left=up_to_date,
                                    right=reduce,
                                    how='inner',
                                    on='ID')
                df_check.to_excel(f'./output/重要名单/核减专属名单/3.{sick_type}两表差异名单.xlsx', index=False)
                df_check = pd.read_excel(f'./output/重要名单/核减专属名单/3.{sick_type}两表差异名单.xlsx')
                # 1.外区+外市+外省
                df_area = (df_check['区县'] != '白云区') & (df_check['地市'] == '广州市') & (df_check['是否排除密接/次密'] != '是') & (
                        df_check['是否追踪到'] != '转出外省') & (df_check['是否核心密接'].str.contains('|'.join(type_type))) & (
                                      df_check['密接类型'] == '密切接触者')
                df_city = (df_check['地市'] != '广州市') & (df_check['是否追踪到'] != '转出外省') & (df_check['是否排除密接/次密'] != '是') & (
                    df_check['是否核心密接'].str.contains('|'.join(type_type))) & (df_check['密接类型'] == '密切接触者')
                df_province = (df_check['是否追踪到'] == '转出外省') & (df_check['是否排除密接/次密'] != '是') & (
                    df_check['是否核心密接'].str.contains('|'.join(type_type))) & (df_check['密接类型'] == '密切接触者')
                # 2.排除身份
                df_except = (df_check['是否排除密接/次密'] == '是') & (df_check['是否核心密接'].str.contains('|'.join(type_type))) & (
                            df_check['密接类型'] == '密切接触者')
                # 3.计算人数
                '''协查到区外人数'''
                df_area_count = len(df_check[df_area])
                df_city_count = len(df_check[df_city])
                df_province_count = len(df_check[df_province])
                df_allarea_count = df_area_count + df_city_count + df_province_count
                # 4.排除身份人数
                df_except_count = len(df_check[df_except])
                # 5.删卡人数(因为有次密身份,无法计算)
                df_delete_count = len(reduce['ID']) - len(df_check['ID'])
                # 6.升降级身份
                if type_type == '是':
                    # 核密降级
                    df_level = (df_check['是否核心密接'] != '是') & (df_check['密接类型'] == '密切接触者')
                    df_level_count = len(df_check[df_level])
                    df_check[df_level].to_excel(f'./output/重要名单/核减专属名单/5.{sick_type}身份降级名单.xlsx')
                elif type_type != '是':
                    df_level = (df_check['是否核心密接'] == '是') & (df_check['密接类型'] == '密切接触者')
                    df_level_count = len(df_check[df_level])
                    df_check[df_level].to_excel(f'./output/重要名单/核减专属名单/5.{sick_type}身份升级名单.xlsx')
                # 7.总排除人数：删卡为排除身份人数
                df_except_allcount = df_except_count + df_delete_count
                # 8.核减总人数
                df_reduce_count = df_allarea_count + df_except_allcount + df_level_count
                # 生成核减名单(协查区外+排除身份+升降级身份)
                df_reduce_dytes = df_area + df_city + df_province + df_except + df_level
                df_check[df_reduce_dytes].to_excel(f'./output/重要名单/核减专属名单/1.{sick_type}核减名单.xlsx', index=False)
                # 7.新增人数
                df_add_dytes = (df_check['区县'] == '白云区') & (df_check['是否排除密接/次密'] != '是') & (
                        df_check['地市'] == '广州市') & (df_check['是否追踪到'] != '转出外省') & (
                                   df_check['是否核心密接'].str.contains('|'.join(type_type))) & (df_check['密接类型'] == '密切接触者')
                df_add_count = len(df_check[df_add_dytes])
                if (df_allarea_count == 0) & (df_level_count == 0):
                    four_replace = f'4.较上一版核减{df_reduce_count}人(排除{sick_type}身份{df_except_allcount}人)；'
                    Original().doc.add_paragraph(four_replace)
                elif (df_allarea_count == 0) & (df_except_allcount == 0):
                    four_replace = f'4.较上一版核减{df_reduce_count}人({sick_type}身份变更{df_level_count}人)；'
                    Original().doc.add_paragraph(four_replace)
                elif (df_level_count == 0) & (df_except_allcount == 0):
                    four_replace = f'4.较上一版核减{df_reduce_count}人(协查到区外{df_allarea_count}人)；'
                    Original().doc.add_paragraph(four_replace)
                elif df_except_allcount == 0:
                    four_replace = f'4.较上一版核减{df_reduce_count}人(协查到区外{df_allarea_count}人,{sick_type}身份变更{df_level_count}人)；'
                    Original().doc.add_paragraph(four_replace)
                elif df_allarea_count == 0:
                    four_replace = f'4.较上一版核减{df_reduce_count}人(排除{sick_type}身份{df_except_allcount}人,{sick_type}身份变更{df_level_count}人)；'
                    Original().doc.add_paragraph(four_replace)
                else:
                    four_replace = f'4.较上一版核减{df_reduce_count}人(协查到区外{df_allarea_count}人,排除{sick_type}身份{df_except_allcount}人)；'
                    Original().doc.add_paragraph(four_replace)
                print(four_replace)

                if df_add_count == 0:
                    five_str = f'5.新增{sick_type}：{last[7:-5]}-{Original().open_file()[6:-5]}无新增。'
                    Original().doc.add_paragraph(five_str)
                    print(five_str)
                    df_check[df_add_dytes].to_excel(f'./output/重要名单/核减专属名单/2.{sick_type}新增名单.xlsx', index=False)
                else:
                    # 生成新增名单
                    df_check[df_add_dytes].to_excel(f'./output/重要名单/核减专属名单/2.{sick_type}新增名单.xlsx', index=False)
                    five_str = f'5.新增{sick_type}：{last[7:-5]}-{Original().open_file()[6:-5]}新增{df_add_count}人。'
                    Original().doc.add_paragraph(five_str)
                    print(five_str)
                print(
                    f'检测出({title})这一版{sick_type}有核减人员,核减人员详情已保存放到/output/重要名单/核减专属名单的文件，删卡只有ID，可以在差异ID查看，（删卡{df_delete_count}人）')
            else:
                print(four_str)
                Original().doc.add_paragraph(four_str)
            print('')
            # 打印名单临时名单(大疫名单)
            if write == '是':
                # 规模疫情专用
                df[focus].to_excel(rf'./output/次要名单/6.{sick_type}集中名单.xlsx', sheet_name=f'{sick_type}集中{focus_count}人',index=False)
                df[home].to_excel(rf'./output/次要名单/7.{sick_type}居家名单.xlsx', sheet_name=f'{sick_type}居家{home_count}人',index=False)
                df[wait].to_excel(rf'./output/重要名单/1.{sick_type}待转运名单.xlsx',sheet_name=f'{sick_type}待转运{wait_count}人', index=False)
                df[track].to_excel(rf'./output/重要名单/2.{sick_type}核实追踪名单.xlsx', f'{sick_type}核实追踪{track_count}人',index=False)
                df[must].to_excel(rf'./output/重要名单/3.{sick_type}涉及我区（我区应管）名单.xlsx', f'涉及我区{sick_type}{must_count}人',index=False)
                # 质控名单
                df['医学观察方式'].fillna('空白', inplace=True)
                wrong_yx = must & (df['医学观察方式'] == '空白') & (df['转归'].str.contains('继续')) & (df['审核时间'] != '空白')
                if len(df[wrong_yx]) != 0:
                    df[wrong_yx].to_excel(rf'./output/质控名单（每次提醒维护后请删除里面的文件）/{sick_type}医学观察方式维护.xlsx', index=False)
                    print(f'检测出{sick_type}数据不闭环，目前在管人数不闭环！！！请检查output/质控名单的{sick_type}医学观察方式维护.xlsx')
                wrong_zg = workable & (df['转归'] == '空白') & (df['审核时间'] != '空白') & (df['审核时间'] != '空白')
                if len(df[wrong_zg]) != 0:
                    df[wrong_zg].to_excel(rf'./output/质控名单（每次提醒维护后请删除里面的文件）/{sick_type}转归维护.xlsx', index=False)
                    print(f'检测出{sick_type}数据不闭环，已落地人数与后面数据不闭环！！！请检查output/质控名单的{sick_type}转归维护.xlsx')
                road_List = ['同德', '松洲', '黄石', '石井', '鹤龙', '龙归', '金沙', '嘉禾', '云城', '三元里', '京溪', '同和', '人和', '均禾', '大源','太和', '白云湖','景泰', '棠景', '永平', '江高', '石门', '新市', '钟落潭']
                wrong_must = must & (df['镇（街道）'].str.contains('|'.join(road_List)) == False) & (
                        df['转归'].str.contains('解除') == False)
                df[wrong_must].to_excel(rf'./output/质控名单（每次提醒维护后请删除里面的文件）/1.质控名单.{sick_type}镇街维护.xlsx', index=False)
                df_write = pd.read_excel(rf'./output/质控名单（每次提醒维护后请删除里面的文件）/1.质控名单.{sick_type}镇街维护.xlsx')
                # 尝试根据目前所在位置写入正确街道
                df_write['目前所处位置'].fillna(df_write['现住址'], inplace=True)
                df_write['目前所处位置'].fillna('不明', inplace=True)
                for j in range(0, len(road_List)):
                    df_write['镇（街道）'].mask(df_write['目前所处位置'].str.contains(f'{str(road_List[j])}'), f'{road_List[j]}街道',inplace=True)
                df_write['镇（街道）'].mask((df_write['镇（街道）'] == '人和街道'), '人和镇', inplace=True)
                df_write['镇（街道）'].mask((df_write['镇（街道）'] == '太和街道'), '太和镇', inplace=True)
                df_write['镇（街道）'].mask((df_write['镇（街道）'] == '钟落潭街道'), '钟落潭镇', inplace=True)
                df_write['镇（街道）'].mask((df_write['镇（街道）'] == '江高街道'), '江高镇', inplace=True)
                df_write.to_excel(rf'./output/质控名单（每次提醒维护后请删除里面的文件）/1.质控名单.{sick_type}镇街维护.xlsx', index=False)
                wrong_must_track = wrong_must & (df[wrong_must]['医学观察方式'].str.contains('待转运')) & (
                        df[wrong_must]['审核时间'] != '空白')
                if len(df[wrong_must_track]) != 0:
                    print(f'检测出{sick_type}待转运有待甄别街道,请检查output/质控名单/1.{sick_type}镇街维护.xlsx')
                wrong_must_wait = wrong_must & (df[wrong_must]['审核时间'] == '空白')
                if len(df[wrong_must_wait]) != 0:
                    print(f'检测出{sick_type}正在落地管控有待甄别街道,请检查output/质控名单/1.{sick_type}镇街维护.xlsx')

            '''
            日常专用
            # 次要名单
            df[active].to_csv(rf'./output/次要名单/1.{sick_type}主动甄别名单.csv', index=False)
            df[province].to_csv(rf'./output/次要名单/2.{sick_type}推送外省名单.csv', index=False)
            df[city].to_csv(rf'./output/次要名单/3.{sick_type}推送外市名单.csv', index=False)
            df[area].to_csv(rf'./output/次要名单/4.{sick_type}推送外区名单.csv', index=False)
            df[workable].to_csv(rf'./output/次要名单/5.已落地{sick_type}核酸名单.csv', index=False)
            df[focus].to_excel(rf'./output/次要名单/6.{sick_type}集中名单.xlsx', index=False)
            df[home].to_excel(rf'./output/次要名单/7.{sick_type}居家名单.xlsx', index=False)
            df[count].to_csv(rf'./output/次要名单/7.累计甄别{sick_type}名单.csv', index=False)

            # 重要名单
            df[wait].to_excel(rf'./output/重要名单/1.{sick_type}待转运名单.xlsx',sheet_name=f'{sick_type}待转运{wait_count}人', index=False)
            df[track].to_excel(rf'./output/重要名单/2.{sick_type}核实追踪名单.xlsx', f'{sick_type}核实追踪{track_count}人',index=False)
            df[must].to_excel(rf'./output/重要名单/3.{sick_type}涉及我区（我区应管）名单.xlsx', f'涉及我区{sick_type}{must_count}人',index=False)

            # 质控名单
            df['医学观察方式'].fillna('空白', inplace=True)
            wrong_yx = must & (df['医学观察方式'] == '空白') & (df['转归'].str.contains('继续')) & (df['审核时间'] != '空白')
            if len(df[wrong_yx]) != 0:
                df[wrong_yx].to_excel(rf'./output/质控名单（每次提醒维护后请删除里面的文件）/{sick_type}医学观察方式维护.xlsx', index=False)
                print(f'检测出{sick_type}数据不闭环，目前在管人数不闭环！！！请检查output/质控名单的{sick_type}医学观察方式维护.xlsx')
            wrong_zg = workable & (df['转归'] == '空白') & (df['审核时间'] != '空白') & (df['审核时间'] != '空白')
            if len(df[wrong_zg]) != 0:
                df[wrong_zg].to_excel(rf'./output/质控名单（每次提醒维护后请删除里面的文件）/{sick_type}转归维护.xlsx', index=False)
                print(f'检测出{sick_type}数据不闭环，已落地人数与后面数据不闭环！！！请检查output/质控名单的{sick_type}转归维护.xlsx')
            road_List = ['同德', '松洲', '黄石', '石井', '鹤龙', '龙归', '金沙', '嘉禾', '云城', '三元里', '京溪', '同和', '人和', '均禾', '大源',
                         '太和', '白云湖',
                         '景泰', '棠景', '永平', '江高', '石门', '新市', '钟落潭']
            wrong_must = must & (df['镇（街道）'].str.contains('|'.join(road_List)) == False)
            df[wrong_must].to_excel(rf'./output/质控名单（每次提醒维护后请删除里面的文件）/1.质控名单.{sick_type}镇街维护.xlsx', index=False)
            # 尝试根据目前所在位置写入正确街道
            df_write = pd.read_excel(rf'./output/质控名单（每次提醒维护后请删除里面的文件）/1.质控名单.{sick_type}镇街维护.xlsx')
            df_write['目前所处位置'].fillna(df_write['现住址'],inplace=True)
            df_write['目前所处位置'].fillna('不明',inplace=True)
            for j in range(0,len(road_List)):
                df_write['镇（街道）'].mask(df_write['目前所处位置'].str.contains(f'{str(road_List[j])}'),f'{road_List[j]}街道',inplace=True)
            df_write['镇（街道）'].mask((df_write['镇（街道）'] == '人和街道'),'人和镇',inplace=True)
            df_write['镇（街道）'].mask((df_write['镇（街道）'] == '太和街道'),'太和镇',inplace=True)
            df_write['镇（街道）'].mask((df_write['镇（街道）'] == '钟落潭街道'),'钟落潭镇',inplace=True)
            df_write['镇（街道）'].mask((df_write['镇（街道）'] == '江高街道'),'江高镇',inplace=True)
            df_write.to_excel(rf'./output/质控名单（每次提醒维护后请删除里面的文件）/1.质控名单.{sick_type}镇街维护.xlsx', index=False)
            wrong_must_track = wrong_must & (df[wrong_must]['医学观察方式'].str.contains('待转运')) & (df[wrong_must]['审核时间'] != '空白')
            if len(df[wrong_must_track]) != 0:
                print(f'检测出{sick_type}待转运有待甄别街道,请检查output/质控名单/1.{sick_type}镇街维护.xlsx')
            wrong_must_wait = wrong_must & (df[wrong_must]['审核时间'] == '空白')
            if len(df[wrong_must_wait]) != 0:
                print(f'检测出{sick_type}正在落地管控有待甄别街道,请检查output/质控名单/1.{sick_type}镇街维护.xlsx')
            '''
            sick_type = '密接'

    def report_simpleusual(self):
        print('\n\n\n')
        #  打印简版标题模板
        print(f'【{title}涉我区密接管控情况】')
        Original().doc.add_paragraph(f'【{title}涉我区密接管控情况】')
        # 读取处理文件
        df = pd.read_csv(f'D:/工作源文件/{Original().open_file()[5:-5]}.csv')
        # 名字类别
        name_List = name.split()
        name_type = df['关联病例'].str.contains('|'.join(name_List))
        # 核心密接+所有密接
        type = ['是', '']
        sick_type = '核心密接'
        for i in range(0, len(type)):
            # 上一版我区应管核密/密接人数/第一点
            # 累计甄别
            count = (df['是否排除密接/次密'] != '是') & (df['是否核心密接'].str.contains('|'.join(type[i]))) & name_type & (df['密接类型'] == '密切接触者')
            # 涉及我区密接、核密
            must = count & (df['是否追踪到'] != '转出外省') & (df['地市'] == '广州市') & (df['区县'] == '白云区')
            must_count = len(df[must])
            if must_count == 0:
                simple_must_str = f'暂无{sick_type}涉及我区；'
            else:
                simple_must_str = f'涉及我区{sick_type}{must_count}人：'
            # 已落地
            workable = (must & (df['审核时间'] != '空白')) + (must & (df['转归'] == '转为确诊'))
            workable_count = len(df[workable])
            if workable_count == 0:
                workable_str = ''
            else:
                workable_str = f'已落地{workable_count}人；'
            # 目前在管
            atpresent = workable & (df['转归'].str.contains('继续'))
            atpresent_count = len(df[atpresent])
            if atpresent_count == 0:
                atpresent_str = ''
            else:
                atpresent_str = f'目前在管{atpresent_count}人'
            # 集中
            focus = atpresent & (df['医学观察方式'].str.contains('集中'))
            focus_count = len(df[focus])
            if focus_count == 0:
                focus_str = ''
            else:
                focus_str = f'集中隔离{focus_count}人，'
            # 待转运
            wait = atpresent & (df['医学观察方式'].str.contains('待转运'))
            wait_count = len(df[wait])
            if wait_count == 0:
                simple_wait_str = ''
            else:
                simple_wait_str = f'待转运{wait_count}人，'
            # 居家
            home = atpresent & (df['医学观察方式'].str.contains('居家'))
            home_count = len(df[home])
            if home_count == 0:
                home_str = ''
            else:
                home_str = f'居家隔离{home_count}人'
            # 医院隔离
            hospital = atpresent & (df['医学观察方式'].str.contains('医院'))
            hospital_count = len(df[hospital])
            if hospital_count == 0:
                hospital_str = ''
            else:
                hospital_str = f'，医院隔离{hospital_count}人，'
            # 解除隔离
            remove = must & (df['转归'].str.contains('解除')) & (df['审核时间'] != '空白')
            remove_count = len(df[remove])
            if remove_count == 0:
                remove_str = ''
            else:
                remove_str = f'解除隔离{remove_count}人；'
            # 正在落地管控
            track = must & (df['审核时间'] == '空白') & (df['转归'] != '转为确诊')
            track_count = len(df[track])
            # 简版
            simple_track_str = f'正在落地管控{track_count}人。'
            if track_count == 0:
                simple_track_str = ''
            else:
                simple_track_str = f'正在落地管控{track_count}人。'
            # 转为确诊
            sun = must & (df['转归'] == '转为确诊')
            sun_count = len(df[sun])
            if sun_count == 0:
                sun_str = ''
            else:
                sun_str = f'转为确诊{sun_count}人；'
            # 简版
            if must_count != workable_count + track_count:
                wrong_str = f'{sick_type}我区应管与已落地+正在落地管控人数不闭环,请检查原因~'
                print(wrong_str)
            if atpresent_count != focus_count + wait_count + home_count + hospital_count:
                wrong_str = f'{sick_type}目前在管人数与(集中隔离+待转运+居家+医院)人数不闭环,请检查原因~'
                print(wrong_str)
            if workable_count != atpresent_count + remove_count + sun_count:
                wrong_str = f'{sick_type}已落地人数与(目前在管+解除隔离+转为确诊)人数不闭环,请检查原因~'
                print(wrong_str)
            simple_second_str = simple_must_str + workable_str + atpresent_str + '【' + focus_str + simple_wait_str + home_str + hospital_str + '】；' + remove_str + sun_str + simple_track_str
            simple_second_str = simple_second_str.replace('，】；','】；')
            # 打印简版
            print(simple_second_str)
            Original().doc.add_paragraph(simple_second_str)
            sick_type = '密接'

    # 4.重点场所报数(目前只删除了次密部分)
    def report_place(self):
        print(f'\n\n\n')
        z = f'【{title}涉我区密接关联重点场所管控情况】（数据来源于省流调系统）'
        print(z)
        # df = pd.read_excel(f'D:/工作源文件/{Original().open_file()[5:]}')
        df = pd.read_csv(f'D:/工作源文件/{Original().open_file()[5:-5]}.csv')
        name_List = name.split()
        name_type = df['关联病例'].str.contains('|'.join(name_List))
        # 大类别
        Indirect_type = df['密接类型'] == '密接的密接'
        Close_type = df['密接类型'] == '密切接触者'
        count = (df['是否排除密接/次密'] != '是') & Close_type & name_type
        counti = (df['是否排除密接/次密'] != '是') & Indirect_type & name_type
        must = count & (df['是否追踪到'] != '转出外省') & (df['地市'] == '广州市') & (df['区县'] == '白云区')
        musti = counti & (df['是否追踪到'] != '转出外省') & (df['地市'] == '广州市') & (df['区县'] == '白云区')

        # 关联重点场所:
        place = []
        # 密接次密
        for col in df[must]['关联重点场所'].unique():
            place.append(col)

        # 我区应管
        for col in df[musti]['关联重点场所'].unique():
            place.append(col)
        # 密接：
        # active = count & (df['创建单位'] == '白云区疾病预防控制中心')
        # 我区应管
        must = count & (df['是否追踪到'] != '转出外省') & (df['地市'] == '广州市') & (df['区县'] == '白云区')
        # 已落地
        workable = must & (df['审核时间'] != '空白')
        workablei = musti & (df['审核时间'] != '空白')
        # 目前在管
        at_present = workable & (df['转归'].str.contains('继续'))
        at_presenti = workablei & (df['转归'].str.contains('继续'))
        # 集中隔离
        focus = at_present & (df['医学观察方式'].str.contains('集中'))
        focusi = at_presenti & (df['医学观察方式'].str.contains('集中'))
        # 待转运
        wait = at_present & (df['医学观察方式'] == '待转运')
        waiti = at_presenti & (df['医学观察方式'] == '待转运')
        # 居家
        home = at_present & (df['医学观察方式'].str.contains('居家'))
        homei = at_presenti & (df['医学观察方式'].str.contains('居家'))
        # 医院隔离
        hospital = at_present & (df['医学观察方式'].str.contains('医院'))
        hospitali = at_presenti & (df['医学观察方式'].str.contains('医院'))
        # 解除观察
        remove = must & (df['转归'].str.contains('解除'))
        removei = musti & (df['转归'].str.contains('解除'))
        # 正在落地管控
        track = must & (df['审核时间'] == '空白')
        tracki = musti & (df['审核时间'] == '空白')
        # 转为阳性
        df['转归'] = df['转归'].str.replace(pat='.*确诊.*', repl='转为确诊', regex=True)
        df['转归'] = df['转归'].str.replace(pat='.*阳性.*', repl='转为确诊', regex=True)
        sun_count = must & (df['转归'] == '转为确诊')
        sun_counti = musti & (df['转归'] == '转为确诊')
        print(f'其中关联重点场所共{len(place)}个，数据如下：')
        # 密接次密
        for i in range(0, len(place)):
            # 累计甄别
            count_place = count & (df['关联重点场所'] == place[i])
            count_placei = counti & (df['关联重点场所'] == place[i])
            countm = len(df[count_place])
            countc = len(df[count_placei])
            if countm == 0:
                coum = ''
            else:
                coum = f'累计甄别{countm}人：'
            if countc == 0:
                couc = ''
            else:
                couc = f'累计甄别{countc}人：'

            # 我区应管
            must_place = must & (df['关联重点场所'] == place[i])
            must_placei = musti & (df['关联重点场所'] == place[i])
            mustm = len(df[must_place])
            mustc = len(df[must_placei])
            if mustm == 0:
                musm = ''
            else:
                musm = f'我区应管{mustm}人；'
            if mustc == 0:
                musc = ''
            else:
                musc = f'我区应管{mustc}人；'

            # 已落地
            workable_place = workable & (df['关联重点场所'] == place[i])
            workable_placei = workablei & (df['关联重点场所'] == place[i])
            workablem = len(df[workable_place])
            workablec = len(df[workable_placei])
            if workablem == 0:
                worm = ''
            else:
                worm = f'已落地{workablem}人；'
            if workablec == 0:
                worc = ''
            else:
                worc = f'已落地{workablec}人；'

            # 目前在管
            present_place = at_present & (df['关联重点场所'] == place[i])
            present_placei = at_presenti & (df['关联重点场所'] == place[i])
            presentm = len(df[present_place])
            presentc = len(df[present_placei])
            if presentm == 0:
                prem = ''
            else:
                prem = f'目前在管{presentm}人'
            if presentc == 0:
                prec = ''
            else:
                prec = f'目前在管{presentc}人'

            # 集中隔离
            focus_place = focus & (df['关联重点场所'] == place[i])
            focus_placei = focusi & (df['关联重点场所'] == place[i])
            focusm = len(df[focus_place])
            focusc = len(df[focus_placei])
            if focusm == 0:
                focm = ''
            else:
                focm = f'集中隔离{focusm}人，'
            if focusc == 0:
                focc = ''
            else:
                focc = f'集中隔离{focusc}人，'

            # 待转运
            wait_place = wait & (df['关联重点场所'] == place[i])
            wait_placei = waiti & (df['关联重点场所'] == place[i])
            waitm = len(df[wait_place])
            waitc = len(df[wait_placei])
            if waitm == 0:
                waim = ''
            else:
                waim = f'待转运{waitm}人，'
            if waitc == 0:
                waic = ''
            else:
                waic = f'待转运{waitc}人'

            # 居家
            home_place = home & (df['关联重点场所'] == place[i])
            home_placei = homei & (df['关联重点场所'] == place[i])
            homem = len(df[home_place])
            homec = len(df[home_placei])
            if homem == 0:
                homm = ''
            else:
                homm = f'居家隔离{homem}人，'
            if homec == 0:
                homc = ''
            else:
                homc = f'居家隔离{homec}人，'

            # 医院隔离
            hospital_place = hospital & (df['关联重点场所'] == place[i])
            hospital_placei = hospitali & (df['关联重点场所'] == place[i])
            hospitalm = len(df[hospital_place])
            hospitalc = len(df[hospital_placei])
            if hospitalm == 0:
                hosm = ''
            else:
                hosm = f'，医院隔离{hospitalm}人；'
            if hospitalc == 0:
                hosc = ''
            else:
                hosc = f'，医院隔离{hospitalc}人；'

            # 解除观察
            remove_place = remove & (df['关联重点场所'] == place[i])
            remove_placei = removei & (df['关联重点场所'] == place[i])
            removem = len(df[remove_place])
            removec = len(df[remove_placei])
            if removem == 0:
                remm = ''
            else:
                remm = f'解除观察{removem}人；'
            if removec == 0:
                remc = ''
            else:
                remc = f'解除观察{removec}人；'

            # 正在落地管控
            track_place = track & (df['关联重点场所'] == place[i])
            track_placei = tracki & (df['关联重点场所'] == place[i])
            trackm = len(df[track_place])
            trackc = len(df[track_placei])
            if trackm == 0:
                tracm = ''
            else:
                tracm = f'正在落地管控{trackm}人；'
            if trackc == 0:
                tracc = ''
            else:
                tracc = f'正在落地管控{trackc}人；'

            # 转为确诊
            sun_place = sun_count & (df['关联重点场所'] == place[i])
            sun_placei = sun_counti & (df['关联重点场所'] == place[i])
            sunm = len(df[sun_place])
            sunc = len(df[sun_placei])
            if sunm == 0:
                sunmm = ''
            else:
                sunmm = f'转为确诊{sunm}人；'
            if sunc == 0:
                suncc = ''
            else:
                suncc = f'转为确诊{sunc}人；'
            # 括号部分：四种情况:

            # print(f'【{place[i]}】')
            # print(f'1.密接：{musm}{worm}{prem}({focm}{homm}{waim}{hosm})；{sunmm}{remm}{tracm}')
            # print(f'2.次密：{musc}{worc}{prec}({focc}{homc}{waic}{hosc})；{suncc}{remc}{tracc}')
            # 在管均0人
            if (presentm == 0) & (presentc == 0):
                print(f'【{place[i]}】')
                print(f'累计甄别密接{countm}人；涉及我区密接{len(df[must_place])}人：{worm}{prem}{sunmm}{remm}{tracm}')
                # print(f'累计甄别次密{countc}人；涉及我区次密{len(df[must_placei])}人：{worc}{prec}{suncc}{remc}{tracc}')

            # 次密在管非0人
            elif (presentm == 0) & (presentc != 0):
                print(f'【{place[i]}】')
                print(f'累计甄别密接{countm}人；涉及我区密接{len(df[must_place])}人：{worm}{prem}{sunmm}{remm}{tracm}')
                # print(f'累计甄别次密{countc}人；涉及我区次密{len(df[must_placei])}人：{worc}{prec}({focc}{homc}{waic}{hosc})；{suncc}{remc}{tracc}')

            # 密接在管非0人
            elif (presentm != 0) & (presentc == 0):
                print(f'【{place[i]}】')
                print(
                    f'累计甄别密接{countm}人；涉及我区密接{len(df[must_place])}人：{worm}{prem}({focm}{homm}{waim}{hosm})；{sunmm}{remm}{tracm}')
                # print(f'累计甄别次密{countc}人；涉及我区次密{len(df[must_placei])}人：{worc}{prec}{suncc}{remc}{tracc}')

            # 均非0人
            else:
                print(f'【{place[i]}】')
                print(
                    f'累计甄别密接{countm}人；涉及我区密接{len(df[must_place])}人：{worm}{prem}({focm}{homm}{waim}{hosm})；{sunmm}{remm}{tracm}')
                # print(f'累计甄别次密{countc}人；涉及我区次密{len(df[must_placei])}人：{worc}{prec}({focc}{homc}{waic}{hosc})；{suncc}{remc}{tracc}')

        print(f'\n\n\n')

    # 5.11表报数（大改，新版11月11日以后）
    def report_new11(self):
        print('正在生成1+1表')
        df = pd.read_csv(f'D:/工作源文件/{Original().open_file()[5:-5]}.csv')
        wb = Workbook()
        wb.save(rf'.\output\1+1表\截止{Original().open_file()[6:-5]}_关联{title[:40]}_1+1表.xlsx')
        wb.close()
        book = load_workbook(rf'.\output\1+1表\截止{Original().open_file()[6:-5]}_关联{title[:40]}_1+1表.xlsx')
        writer = pd.ExcelWriter(rf'.\output\1+1表\截止{Original().open_file()[6:-5]}_关联{title[:40]}_1+1表.xlsx')
        writer.book = book
        # 开始统计:关联病例
        name_List = name.split()
        name_type = df['关联病例'].str.contains('|'.join(name_List))
        # 质控街道
        countall = (df['是否排除密接/次密'] != '是') & name_type
        mustall = countall & (df['是否追踪到'] != '转出外省') & (df['地市'] == '广州市') & (df['区县'] == '白云区')
        road_List = ['同德', '松洲', '黄石', '石井', '鹤龙', '龙归', '金沙', '嘉禾', '云城', '三元里', '京溪', '同和', '人和', '均禾', '大源', '太和',
                     '白云湖','景泰', '棠景', '永平', '江高', '石门', '新市', '钟落']
        wrong_must = mustall & (df['镇（街道）'].str.contains('|'.join(road_List)) == False)
        df['镇（街道）'][wrong_must] = '核实'

        # 首先生成6个名单,再生成统计表
        '''所有密接未管控，所有密接未转运，非核心密接未管控，非核心密接未转运，核心密接未管控,核心密接未转运'''

        df['是否超6小时未管控'] = pd.to_datetime(df['录入时间'])
        df['是否超6小时未转运'] = df['审核时间']
        now_time = str(datetime.datetime.now())
        df['当前日期'] = now_time[:-10]
        df['录入时间'] = pd.to_datetime(df['录入时间'])
        df['当前日期'] = pd.to_datetime(df['当前日期'])
        # 大类别
        type = ['','否','是']  # 核心密接或所有密接
        sick_type = ['所有密接','非核心密接','核心密接']  # 对应sick_type
        wait_sheet = []
        track_sheet = []
        for i in range(0,len(type)):
            # 累计甄别
            count = (df['是否排除密接/次密'] != '是') & (df['是否核心密接'].str.contains('|'.join(type[i]))) & name_type & (df['密接类型'] == '密切接触者')
            # 我区应管
            must = count & (df['是否追踪到'] != '转出外省') & (df['地市'] == '广州市') & (df['区县'] == '白云区')
            # 已落地
            workable = (must & (df['审核时间'] != '空白')) + (must & (df['转归'] == '转为确诊'))
            # 目前在管
            atpresent = workable & (df['转归'].str.contains('继续'))
            # 待转运
            wait = atpresent & (df['医学观察方式'].str.contains('待转运'))
            wait_count = len(df[wait])
            # 正在落地管控
            track = must & (df['审核时间'] == '空白') & (df['转归'] != '转为确诊')
            track_count = len(df[track])
            df[wait].to_excel(writer, sheet_name=f'{sick_type[i]}未转运{wait_count}人',index=False)
            df[track].to_excel(writer, sheet_name=f'{sick_type[i]}未管控{track_count}人',index=False)
            wait_sheet.append(f'{sick_type[i]}未转运{wait_count}人')
            track_sheet.append(f'{sick_type[i]}未管控{track_count}人')
        writer.save()
        writer.close()
        # 开始打印两个统计表
        wb = load_workbook(filename=rf'.\output\1+1表\截止{Original().open_file()[6:-5]}_关联{title[:40]}_1+1表.xlsx')
        del wb['Sheet']

        # 两个统计表
        wb_Uncontrol_count = wb.create_sheet('未管控未转运人员统计')
        '''1.待管控未转运人员名单(合算)'''
        # 第一个表
        # 所有街道
        row7_30 = ['同德街道', '松洲街道', '黄石街道', '石井街道', '鹤龙街道', '龙归街道', '金沙街道', '嘉禾街道', '云城街道', '三元里街道', '京溪街道', '同和街道',
                   '人和镇',
                   '均禾街道', '大源街道', '太和镇', '白云湖街道', '景泰街道', '棠景街道', '永平街道', '江高镇', '石门街道', '新市街道', '钟落潭镇', '街道待核实']
        # 填充未管控未转运人员统计
        wb_Uncontrol_count['A3'] = '序号'
        wb_Uncontrol_count['B3'] = '镇街'
        wb_Uncontrol_count['C3'] = '所有密接'
        wb_Uncontrol_count['G3'] = '非核心密接'
        wb_Uncontrol_count['I3'] = '核心密接'
        wb_Uncontrol_count['A6'] = '总计'
        # 批量填充
        list4 = ['未管控总数', '超6小时未管控', '未转运总数', '超6小时未转运', '未管控人数', '未转运人数', '未管控人数', '未转运人数']
        for i in range(0,len(list4)):
            wb_Uncontrol_count.cell(row=4, column=i+3).value = list4[i]
        # 打印序号及镇街
        for i in range(0, len(row7_30)):
            wb_Uncontrol_count[f'A{i + 7}'] = i + 1
            wb_Uncontrol_count[f'B{i + 7}'] = row7_30[i]

        road_List = ['同德', '松洲', '黄石', '石井', '鹤龙', '龙归', '金沙', '嘉禾', '云城', '三元', '京溪', '同和', '人和', '均禾', '大源', '太和',
                     '白云湖',
                     '景泰', '棠景', '永平', '江高', '石门', '新市', '钟落', '核实']
        # 所有密接\非核\核密循环
        wait_road = ['E','H','J']
        track_road = ['C','G','I']
        for i in range(0,len(wait_sheet)):
            df_wait = pd.read_excel(rf'.\output\1+1表\截止{Original().open_file()[6:-5]}_关联{title[:40]}_1+1表.xlsx',sheet_name=f'{wait_sheet[i]}')
            df_track = pd.read_excel(rf'.\output\1+1表\截止{Original().open_file()[6:-5]}_关联{title[:40]}_1+1表.xlsx',sheet_name=f'{track_sheet[i]}')
            for j in range(0,len(road_List)):
                wait = (df_wait['镇（街道）'].str.contains(f'{road_List[j]}'))
                track = (df_track['镇（街道）'].str.contains(f'{road_List[j]}'))
                wait_all = len(df_wait[wait])
                track_all = len(df_track[track])
                wb_Uncontrol_count[f'{wait_road[i]}{j+7}'] = wait_all
                wb_Uncontrol_count[f'{track_road[i]}{j+7}'] = track_all
        # 批量统计
        number = ['C','D','E','F','G','H','I','J']
        for i in range(0,len(number)):
            wb_Uncontrol_count[f'{number[i]}6'] = f'=SUM({number[i]}7:{number[i]}31)'

        '''设计样式'''
        # 1.合拼单元格！
        wb_Uncontrol_count.merge_cells('A3:A5')
        wb_Uncontrol_count.merge_cells('B3:B5')
        wb_Uncontrol_count.merge_cells('C3:F3')
        wb_Uncontrol_count.merge_cells('G3:H3')
        wb_Uncontrol_count.merge_cells('I3:J3')
        wb_Uncontrol_count.merge_cells('A6:B6')
        wb_Uncontrol_count.merge_cells('C4:C5')
        wb_Uncontrol_count.merge_cells('D4:D5')
        wb_Uncontrol_count.merge_cells('E4:E5')
        wb_Uncontrol_count.merge_cells('F4:F5')
        wb_Uncontrol_count.merge_cells('G4:G5')
        wb_Uncontrol_count.merge_cells('H4:H5')
        wb_Uncontrol_count.merge_cells('I4:I5')
        wb_Uncontrol_count.merge_cells('J4:J5')
        # 2.自动列宽！！
        lks = []  # 英文变量太费劲，用汉语首字拼音代替
        for i in range(1, wb_Uncontrol_count.max_column + 1):  # 每列循环
            lk = 1  # 定义初始列宽，并在每个行循环完成后重置
            for j in range(1, wb_Uncontrol_count.max_row + 1):  # 每行循环
                sz = wb_Uncontrol_count.cell(row=j, column=i).value  # 每个单元格内容
                if isinstance(sz, str):  # 中文占用多个字节，需要分开处理
                    lk1 = len(sz.encode('gbk'))  # gbk解码一个中文两字节，utf-8一个中文三字节，gbk合适
                else:
                    lk1 = len(str(sz))
                if lk < lk1:
                    lk = lk1  # 借助每行循环将最大值存入lk中
            lks.append(lk)  # 将每列最大宽度加入列表。（犯了一个错，用lks = lks.append(lk)报错，append会修改列表变量，返回值none，而none不能继续用append方法）
        for i in range(1, wb_Uncontrol_count.max_column + 1):
            k = get_column_letter(i)  # 将数字转化为列名,26个字母以内也可以用[chr(i).upper() for i in range(97, 123)]，不用导入模块
            wb_Uncontrol_count.column_dimensions[k].width = lks[i - 1] + 2  # 设置列宽，一般加两个字节宽度，可以根据实际情况灵活调整
        # 3.居中对齐！！！
        max_rows = wb_Uncontrol_count.max_row  # 获取最大行
        max_columns = wb_Uncontrol_count.max_column  # 获取最大列
        align = Alignment(horizontal='center', vertical='center')
        # openpyxl的下标从1开始
        for i in range(1, max_rows + 1):
            for j in range(1, max_columns + 1):
                wb_Uncontrol_count.cell(i, j).alignment = align
        # 4.所有框线！！！！
        side = Side(style="thin")
        bord = Border(top=side,  # 上
            bottom=side,  # 下
            left=side,  # 左
            right=side,  # 右
            diagonal=side  # 对角线
        )
        for r in wb_Uncontrol_count:
            for c in r:
                c.alignment = Alignment(wrapText=True, horizontal='center', vertical='center')
                c.border = bord
                c.font = Font(bold=True)
        # 5.设置颜色
        file = PatternFill("solid", fgColor="BDD7EE")
        for i in range(3, 6):
            for j in range(1, 11):
                wb_Uncontrol_count.cell(row=i, column=j).fill = file

        # 6.将非0的数字着重显示
        cell = wb_Uncontrol_count['C7:J31']
        yellow_fill = PatternFill("solid", start_color='FFFF00')
        for row in cell:  # 遍历每一行的单元格
            for column in row:  # 遍历每一列的单元格
                if column.value == 0:
                    column.font = Font(bold=False)
                else:
                    column.font = Font(bold=True)
                    column.fill = yellow_fill
        # 最后设置抬头
        wb_Uncontrol_count.merge_cells('A1:J2')
        wb_Uncontrol_count['A1'] =  f'{title}涉我区密接次密待转运及待管控情况\n{Original().open_file()[6:-5]}数据'
        wb_Uncontrol_count['A1'].font = Font(name="微软雅黑", size=15, bold=True)
        wb_Uncontrol_count.row_dimensions[1].height=22
        wb_Uncontrol_count.row_dimensions[2].height=22

        # 原1+1表大类别（原1+1表）
        wb_allcount = wb.create_sheet('总统计表')
        df = pd.read_csv(f'D:/工作源文件/{Original().open_file()[5:-5]}.csv')
        Close_type = df['密接类型'] == '密切接触者'
        name_List = name.split()
        name_type = df['关联病例'].str.contains('|'.join(name_List))
        # 密接待核实街道
        countm = (df['是否排除密接/次密'] != '是') & Close_type & name_type & (df['是否追踪到'] != '转出外省')
        mustm = countm & (df['是否追踪到'] != '转出外省') & (df['地市'] == '广州市') & (df['区县'] == '白云区')
        road_List = ['同德', '松洲', '黄石', '石井', '鹤龙', '龙归', '金沙', '嘉禾', '云城', '三元里', '京溪', '同和', '人和', '均禾', '大源', '太和',
                     '白云湖',
                     '景泰', '棠景', '永平', '江高', '石门', '新市', '钟落']
        wrong_mustm = mustm & (df['镇（街道）'].str.contains('|'.join(road_List)) == False)
        df['镇（街道）'][wrong_mustm] = '核实'
        road_List = ['同德', '松洲', '黄石', '石井', '鹤龙', '龙归', '金沙', '嘉禾', '云城', '三元', '京溪', '同和', '人和', '均禾', '大源', '太和',
                     '白云湖',
                     '景泰', '棠景', '永平', '江高', '石门', '新市', '钟落', '核实']
        wb_allcount['C6'] = '=D6+H6+I6'
        # 密接循环
        for i in range(0, len(road_List)):
            # 街道
            road = mustm & df['镇（街道）'].str.contains(f'{road_List[i]}')  # & (df['转归'].str.contains('解除') == False)
            #  待核实
            track_11 = road & (df['审核时间'] == '空白')
            wb_allcount[f'D{i + 7}'] = len(df[track_11])
            # 待转运
            wait_11 = road & (df['医学观察方式'] == '待转运') & (df['转归'].str.contains('继续')) & (df['审核时间'] != '空白')
            wb_allcount[f'E{i + 7}'] = len(df[wait_11])
            # 集中
            focus_11 = road & df['医学观察方式'].str.contains('集中') & (df['转归'].str.contains('继续')) & (
                    df['审核时间'] != '空白')
            wb_allcount[f'F{i + 7}'] = len(df[focus_11])
            # # 居家
            home_11 = road & df['医学观察方式'].str.contains('居家') & (df['转归'].str.contains('继续')) & (
                    df['审核时间'] != '空白')
            wb_allcount[f'G{i + 7}'] = len(df[home_11])
            # # 小计
            small_count_11 = road & (df['医学观察方式'].str.contains('集中|居家|待转运')) & (df['转归'].str.contains('继续')) & (
                    df['审核时间'] != '空白')
            wb_allcount[f'H{i + 7}'] = len(df[small_count_11])
            # 解除观察
            remove_11 = mustm & df['镇（街道）'].str.contains(f'{road_List[i]}') & (df['转归'].str.contains('解除')) & (
                    df['审核时间'] != '空白')
            wb_allcount[f'I{i + 7}'] = len(df[remove_11])
            wb_allcount[f'C{i+7}'] = f'=D{i+7}+H{i+7}+I{i+7}'



        letter = ['D', 'E', 'F', 'G', 'H', 'I']
        for i in range(0, len(letter)):
            wb_allcount[f'{letter[i]}6'] = f'=SUM({letter[i]}7:{letter[i]}31)'

        #
        wb_allcount['E5'] = '未转运'
        wb_allcount['F5'] = '集中隔离'
        wb_allcount['G5'] = '居家隔离'
        wb_allcount['H5'] = '小计'

        # 先合拼/居中顶端对齐单元格/样式设置/边框样式，可选{'thick', 'dashDotDot', 'dashed', 'medium', 'mediumDashDotDot', 'mediumDashed', 'slantDashDot', 'dotted', 'double', 'thin', 'hair', 'mediumDashDot', 'dashDot'}
        side = Side(style="thin")
        bord = Border(
            top=side,  # 上
            bottom=side,  # 下
            left=side,  # 左
            right=side,  # 右
            diagonal=side  # 对角线
        )
        area = wb_allcount['A3:A5']
        for i in area:
            for j in i:
                j.border = bord
                j.font = Font(bold=True)

        wb_allcount.merge_cells('A3:A5')
        wb_allcount['A3'] = '序号'
        wb_allcount['A3'].alignment = Alignment(horizontal='center', vertical="top")
        area = wb_allcount['A3:A5']
        for i in area:
            for j in i:
                j.border = bord
                j.font = Font(bold=True)

        wb_allcount.merge_cells('B3:B5')
        wb_allcount['B3'] = '镇街'
        wb_allcount['B3'].alignment = Alignment(horizontal='center', vertical="top")
        wb_allcount['B3'].border = bord
        area = wb_allcount['B3:B5']
        for i in area:
            for j in i:
                j.border = bord
                j.font = Font(bold=True)

        wb_allcount.merge_cells('C3:C5')
        wb_allcount['C3'] = '合计'
        wb_allcount['C3'].alignment = Alignment(horizontal='center', vertical="top")
        area = wb_allcount['C3:C5']
        for i in area:
            for j in i:
                j.border = bord
                j.font = Font(bold=True)

        wb_allcount.merge_cells('D3:I3')
        wb_allcount['D3'] = '密接'
        wb_allcount['D3'].alignment = Alignment(horizontal='center', vertical="center")
        area = wb_allcount['D3:I5']
        for i in area:
            for j in i:
                j.border = bord
                j.font = Font(bold=True)


        wb_allcount.merge_cells('D4:D5')
        wb_allcount['D4'] = '未管控'
        wb_allcount['D4'].alignment = Alignment(horizontal='center', vertical="top")
        area = wb_allcount['D4:D5']
        for i in area:
            for j in i:
                j.border = bord
                j.font = Font(bold=True)

        wb_allcount.merge_cells('I4:I5')
        wb_allcount['I4'] = '解除'
        wb_allcount['I4'].alignment = Alignment(horizontal='center', vertical="top")
        area = wb_allcount['I4:I5']
        for i in area:
            for j in i:
                j.border = bord
                j.font = Font(bold=True)

        wb_allcount.merge_cells('E4:H4')
        wb_allcount['E4'] = '已管控'
        wb_allcount['E4'].alignment = Alignment(horizontal='center', vertical="center")
        area = wb_allcount['E3:H4']
        for i in area:
            for j in i:
                j.border = bord
                j.font = Font(bold=True)


        wb_allcount.merge_cells('A6:B6')
        wb_allcount['A6'] = '合计'
        wb_allcount['A6'].alignment = Alignment(horizontal='center', vertical="center")
        area = wb_allcount['A6:B6']
        for i in area:
            for j in i:
                j.border = bord
                j.font = Font(bold=True)


        row7_30 = ['同德街道', '松洲街道', '黄石街道', '石井街道', '鹤龙街道', '龙归街道', '金沙街道', '嘉禾街道', '云城街道', '三元里街道', '京溪街道', '同和街道',
                   '人和镇',
                   '均禾街道', '大源街道', '太和镇', '白云湖街道', '景泰街道', '棠景街道', '永平街道', '江高镇', '石门街道', '新市街道', '钟落潭镇', '街道待核实']
        # 打印序号及镇街
        for i in range(0, len(row7_30)):
            wb_allcount[f'A{i + 7}'] = i + 1
            wb_allcount[f'A{i + 7}'].alignment = Alignment(horizontal='center', vertical="center")
            wb_allcount[f'A{i + 7}'].border = bord
            wb_allcount[f'B{i + 7}'] = row7_30[i]

        file = PatternFill("solid", fgColor="BDD7EE")
        for i in range(3, 6):
            for j in range(1, 10):
                wb_allcount.cell(row=i, column=j).fill = file

        list = range(0, 26)
        for i in list:
            wb_allcount.cell(row=i + 6, column=3).fill = file






        # 调整顺序
        wb.move_sheet(f"{wait_sheet[2]}", -2)
        wb.move_sheet(f"{track_sheet[2]}", -2)
        wb.move_sheet("未管控未转运人员统计", -6)
        wb.move_sheet("总统计表", -6)
        # 先删除非核密名单,没用
        del wb[f"{wait_sheet[1]}"]
        del wb[f"{track_sheet[1]}"]
        wb.save(rf'.\output\1+1表\截止{Original().open_file()[6:-5]}_关联{title[:40]}_1+1表.xlsx')
        print(f'截止{Original().open_file()[6:-5]}_关联{title[:40]}_1+1表.xlsx已生成')

    # 6.生成名单(所有名单:包含质控)
    def report_roster(self):
        # 先计算密接类型
        sick_type = '密接'
        # 读取处理文件
        # df = pd.read_csv(f'D:/工作源文件/{Original().open_file()[5:-5]}.csv')
        df = pd.read_excel(f'./{Original().open_file()}')
        # 名字类别
        name_List = name.split()
        name_type = df['关联病例'].str.contains('|'.join(name_List))
        print('程序正在进行数据筛选，筛选完成后才会生成名单.....')
        # 密接类型
        type = ['密切接触者', '密接的密接']
        for i in range(0, len(type)):
            # 插入:统计第四点:先统计第四点
            type_type = type[i]
            # 上一版日期和上一版数据
            path2 = r"./last"
            df_last = os.listdir(path2)
            last = os.path.join(path2, df_last[-1])
            df_lastexcel = pd.read_excel(f'{last}')
            # 上一版我区应管密接人数
            last_name = name.split()
            last_name_type = df_lastexcel['关联病例'].str.contains('|'.join(last_name))
            mi = (df_lastexcel['是否追踪到'] != '转出外省') & (df_lastexcel['地市'] == '广州市') & (
                    df_lastexcel['区县'] == '白云区') & (
                         df_lastexcel['是否排除密接/次密'] != '是') & (df_lastexcel['密接类型'] == type[i]) & last_name_type

            '''筛选数据'''
            # 累计甄别
            count_last = len(df_lastexcel[mi])
            count = (df['是否排除密接/次密'] != '是') & (df['密接类型'] == type[i]) & name_type
            # 我区主动甄别
            active = count & (df['创建单位'] == '白云区疾病预防控制中心')
            active_count = len(df[active])
            # 推送外省管控人数
            province = active & (df['是否追踪到'] == '转出外省')
            province_count = len(df[province])
            # 推送外市管控人数
            city = active & (df['是否追踪到'] != '转出外省') & (df['地市'] != '广州市')
            city_count = len(df[city])
            # 推送外区管控人数
            area = active & (df['是否追踪到'] != '转出外省') & (df['地市'] == '广州市') & (df['区县'] != '白云区')
            area_count = len(df[area])
            # 归属白云区
            local = active & (df['是否追踪到'] != '转出外省') & (df['地市'] == '广州市') & (df['区县'] == '白云区')
            local_count = len(df[local])
            # 区外推送直接计算
            pushme_count = len(df[count]) - len(df[active])
            # 外区推送我区,我区推送外省(求Y的值) :区外推送X人（推送外省'Y'人）
            out = count & (df['创建单位'] != '白云区疾病预防控制中心') & (df["是否追踪到"] == "转出外省")
            out_count = len(df[out])
            # 第二点
            # 涉及我区应管
            must = count & (df['是否追踪到'] != '转出外省') & (df['地市'] == '广州市') & (df['区县'] == '白云区')
            must_count = len(df[must])
            # 已落地
            workable = (must & (df['审核时间'] != '空白')) + (must & (df['转归'] == '转为确诊'))
            workable_count = len(df[workable])
            # 目前在管
            atpresent = workable & (df['转归'].str.contains('继续'))
            atpresent_count = len(df[atpresent])
            # 集中
            focus = atpresent & (df['医学观察方式'].str.contains('集中'))
            focus_count = len(df[focus])
            # 待转运
            wait = atpresent & (df['医学观察方式'].str.contains('待转运'))
            wait_count = len(df[wait])
            # 居家
            home = atpresent & (df['医学观察方式'].str.contains('居家'))
            home_count = len(df[home])
            # 医院隔离
            hospital = atpresent & (df['医学观察方式'].str.contains('医院'))
            hospital_count = len(df[hospital])
            # 解除隔离
            remove = must & (df['转归'].str.contains('解除')) & (df['审核时间'] != '空白')
            remove_count = len(df[remove])
            # 正在落地管控
            track = must & (df['审核时间'] == '空白') & (df['转归'] != '转为确诊')
            track_count = len(df[track])
            # 转为确诊
            sun = must & (df['转归'] == '转为确诊')
            sun_count = len(df[sun])
            # 第三点
            # 累计已推送区外管控人数（直接计算）
            allpush_count = province_count + city_count + area_count + out_count
            # 第四点(字面新增)
            add_count = must_count - count_last
            print(f'{sick_type}数据筛选完成,正在生成名单')
            time.sleep(2)
            print('正在进行两表比对,生成核减专属名单,并生成累计新增名单及实际新增名单,该部分处理时间较长,请耐心等待....')
            # 1-生成核减需求名单
            # 找出核减的ID
            reduce = pd.concat([df_lastexcel[mi]['ID'], df[must]['ID']]).drop_duplicates(keep=False)
            # 保存差异的ID.excel
            reduce.to_csv(f'./output/重要名单/核减专属名单/6.两表{sick_type}差异ID.csv', index=False)
            # 读取这个IDExcel
            reduce = pd.read_csv(f'测试代码/output/重要名单/核减专属名单/6.两表{sick_type}差异ID.csv')
            # 读取原始input的excel
            up_to_date = pd.read_excel(f'./{Original().open_file()}')
            # 两表的差异
            df_check = pd.merge(left=up_to_date,
                                right=reduce,
                                how='inner',
                                on='ID')
            df_check.to_csv(f'./output/重要名单/核减专属名单/7.{sick_type}两表差异名单.csv', index=False)
            df_check = pd.read_csv(f'测试代码/output/重要名单/核减专属名单/7.{sick_type}两表差异名单.csv')
            # 1.外区+外市+外省
            df_area = (df_check['区县'] != '白云区') & (df_check['地市'] == '广州市') & (df_check['是否排除密接/次密'] == '否') & (
                        df_check['是否追踪到'] != '转出外省') & (df_check['密接类型'] == type_type)
            df_city = (df_check['地市'] != '广州市') & (df_check['是否追踪到'] != '转出外省') & (df_check['是否排除密接/次密'] == '否') & (
                        df_check['密接类型'] == type_type)
            df_province = (df_check['是否追踪到'] == '转出外省') & (df_check['是否排除密接/次密'] == '否') & (
                        df_check['密接类型'] == type_type)
            # 2.排除身份
            df_except = (df_check['是否排除密接/次密'] == '是') & (df_check['密接类型'] == type_type)
            df_check[df_except].to_excel(f'./output/重要名单/核减专属名单/3.排除{sick_type}身份名单.xlsx', index=False)
            # 3.计算人数
            '''协查到区外人数'''
            df_area_count = len(df_check[df_area])
            df_city_count = len(df_check[df_city])
            df_province_count = len(df_check[df_province])
            df_allarea_count = df_area_count + df_city_count + df_province_count
            # 4.排除身份人数
            df_except_count = len(df_check[df_except])
            # 5.删卡人数
            df_delete_count = len(reduce['ID']) - len(df_check['ID'])
            # 6.升降级身份
            if type_type == '密切接触者':
                # 密接降级
                df_level = (df_check['密接类型'] == '密接的密接')
                df_level_count = len(df_check[df_level])
                df_check[df_level].to_excel(f'./output/重要名单/核减专属名单/4.{sick_type}身份变更名单.xlsx')
            elif type_type == '密接的密接':
                df_level = (df_check['密接类型'] == '密切接触者')
                df_level_count = len(df_check[df_level])
                df_check[df_level].to_excel(f'./output/重要名单/核减专属名单/4.{sick_type}身份变更名单.xlsx')
            # 7.总排除人数：删卡为排除身份人数
            df_except_allcount = df_except_count + df_delete_count
            # 8.核减总人数
            df_reduce_count = df_allarea_count + df_except_allcount + df_level_count
            # 生成核减名单(协查区外+排除身份+升降级身份)
            df_reduce_dytes = df_area + df_city + df_province + df_except + df_level
            df_check[df_reduce_dytes].to_excel(f'./output/重要名单/核减专属名单/1.{sick_type}核减名单.xlsx', index=False)
            # 7.新增人数
            df_add_dytes = (df_check['区县'] == '白云区') & (df_check['是否排除密接/次密'] != '是') & (
                    df_check['地市'] == '广州市') & (df_check['是否追踪到'] != '转出外省') & (df_check['密接类型'] == type_type)
            df_add_count = len(df_check[df_add_dytes])
            df_check[df_add_dytes].to_excel(f'./output/重要名单/核减专属名单/2.{sick_type}新增名单.xlsx', index=False)
            print(
                f'累计核减{df_reduce_count}人(协查到区外{df_allarea_count}人,排除{sick_type}身份{df_except_allcount}人,{sick_type}身份变更{df_level_count}人)。')
            print(f'累计新增{sick_type}：{last[7:-5]}-{Original().open_file()[6:-5]}新增{df_add_count}人。')
            print(
                f'实际新增{sick_type}：{last[7:-5]}-{Original().open_file()[6:-5]}新增{add_count}人。如遇核减+累计新增不相等的情况,以实际新增人数为准')
            print(f'已生成{sick_type}核减类型名单')
            print(f'正在生成{sick_type}次要名单')
            # 打印名单
            # 临时名单

            # 生成次要名单
            df[active].to_csv(rf'./output/次要名单/1.{sick_type}主动甄别名单.csv', index=False)
            print(f'已生成{sick_type}主动甄别名单')
            df[province].to_csv(rf'./output/次要名单/2.{sick_type}推送外省名单.csv', index=False)
            print(f'已生成{sick_type}推送外省名单')
            df[city].to_csv(rf'./output/次要名单/3.{sick_type}推送外市名单.csv', index=False)
            print(f'已生成{sick_type}推送外市名单')
            df[area].to_csv(rf'./output/次要名单/4.{sick_type}推送外区名单.csv', index=False)
            print(f'已生成{sick_type}推送外区名单')
            df[workable].to_csv(rf'./output/次要名单/5.已落地{sick_type}名单.csv', index=False)
            print(f'已生成{sick_type}已落地名单')
            df[focus].to_excel(rf'./output/次要名单/6.{sick_type}集中名单.xlsx', index=False)
            print(f'已生成{sick_type}集中名单')
            df[home].to_excel(rf'./output/次要名单/7.{sick_type}居家名单.xlsx', index=False)
            print(f'已生成{sick_type}居家名单')
            df[count].to_csv(rf'./output/次要名单/8.累计甄别{sick_type}名单.csv', index=False)
            print(f'已生成{sick_type}累计甄别名单')
            # 重要名单
            if sick_type != '次密':
                df[wait].to_excel(rf'./output/重要名单/1.{sick_type}待转运名单.xlsx', index=False)
                print(f'已生成{sick_type}待转运名单')
                core_wait = wait & (df['是否核心密接'] == '是')
                df[core_wait].to_excel(rf'./output/重要名单/1.核心{sick_type}待转运名单.xlsx', index=False)
                core_track = track & (df['是否核心密接'] == '是')
                df[core_track].to_excel(rf'./output/重要名单/1.核心{sick_type}待管控名单.xlsx', index=False)
                print(f'已生成{sick_type}核心密接待管控名单')
            df[track].to_excel(rf'./output/重要名单/2.{sick_type}核实追踪名单.xlsx', index=False)
            print(f'已生成{sick_type}待核实名单')
            df[must].to_excel(rf'./output/重要名单/3.{sick_type}涉及我区（我区应管）名单.xlsx', index=False)
            print(f'已生成涉及我区{sick_type}名单')

            # 质控名单
            # 1.医学观察方式维护
            df['医学观察方式'].fillna('空白', inplace=True)
            wrong_yx = must & (df['医学观察方式'] == '空白') & (df['转归'].str.contains('继续')) & (df['审核时间'] != '空白')
            if len(df[wrong_yx]) != 0:
                df[wrong_yx].to_excel(rf'./output/质控名单（每次提醒维护后请删除里面的文件）/{sick_type}医学观察方式维护.xlsx', index=False)
                print(f'已生成{sick_type}医学观察方式维护名单')
            # 2.转归维护
            wrong_zg = workable & (df['转归'] == '空白') & (df['审核时间'] != '空白') & (df['审核时间'] != '空白')
            if len(df[wrong_zg]) != 0:
                df[wrong_zg].to_excel(rf'./output/质控名单（每次提醒维护后请删除里面的文件）/{sick_type}转归维护.xlsx', index=False)
                print(f'已生成{sick_type}转归维护名单')
            # 3.1街道维护
            road_List = ['同德', '松洲', '黄石', '石井', '鹤龙', '龙归', '金沙', '嘉禾', '云城', '三元里', '京溪', '同和', '人和', '均禾', '大源',
                         '太和', '白云湖',
                         '景泰', '棠景', '永平', '江高', '石门', '新市', '钟落潭']
            wrong_must = must & (df['镇（街道）'].str.contains('|'.join(road_List)) == False)
            df[wrong_must].to_excel(rf'./output/质控名单（每次提醒维护后请删除里面的文件）/1.质控名单.{sick_type}镇街维护.xlsx', index=False)
            print(f'已生成{sick_type}街道维护名单')
            # 3.2尝试根据目前所在位置写入正确街道
            df_write = pd.read_excel(rf'./output/质控名单（每次提醒维护后请删除里面的文件）/1.质控名单.{sick_type}镇街维护.xlsx')
            print(f'正在尝试根据目前所在位置写入街道')
            df_write['目前所处位置'].fillna(df_write['现住址'], inplace=True)
            df_write['目前所处位置'].fillna('不明', inplace=True)
            for j in range(0, len(road_List)):
                df_write['镇（街道）'].mask(df_write['目前所处位置'].str.contains(f'{str(road_List[j])}'), f'{road_List[j]}街道',
                                       inplace=True)
            df_write['镇（街道）'].mask((df_write['镇（街道）'] == '人和街道'), '人和镇', inplace=True)
            df_write['镇（街道）'].mask((df_write['镇（街道）'] == '太和街道'), '太和镇', inplace=True)
            df_write['镇（街道）'].mask((df_write['镇（街道）'] == '钟落潭街道'), '钟落潭镇', inplace=True)
            df_write['镇（街道）'].mask((df_write['镇（街道）'] == '江高街道'), '江高镇', inplace=True)
            df_write.to_excel(rf'./output/质控名单（每次提醒维护后请删除里面的文件）/1.质控名单.{sick_type}镇街维护.xlsx', index=False)
            wrong_must_track = wrong_must & (df[wrong_must]['医学观察方式'].str.contains('待转运')) & (
                        df[wrong_must]['审核时间'] != '空白')
            print(f'填写街道完毕，无法识别的地址无法填入，已重新生成{sick_type}质控名单')
            if len(df[wrong_must_track]) != 0:
                print(f'本次已检测出{sick_type}待转运有待甄别街道，需要及时维护')
            # 4.居家隔离原因为空白,但是医学观察方式为居家改成待转运
            home_check = atpresent & (df['医学观察方式'].str.contains('居家')) & (df['居家隔离原因'].isnull())
            if sick_type != '次密':
                df[home_check].to_excel(rf'./output/质控名单（每次提醒维护后请删除里面的文件）/2.质控名单.{sick_type}居家维护成待转运.xlsx', index=False)
                print(f'{sick_type}居家维护成待转运名单已生成')
            job = must & (df['职业'].isnull())
            df[job].to_excel(rf'./output/质控名单（每次提醒维护后请删除里面的文件）/3.质控名单.{sick_type}职业为空的维护.xlsx', index=False)
            print(f'{sick_type}职业为空的维护名单已生成')
            job_list = ['医护人员', '学生', '老师']
            check_job = must & df['职业'].str.contains('|'.join(job_list)) & (df['工作单位'].isnull())
            df[check_job].to_excel(rf'./output/质控名单（每次提醒维护后请删除里面的文件）/4.质控名单.{sick_type}特定职业工作单位为空维护.xlsx', index=False)
            print(f'{sick_type}特定职业工作单位为空维护名单已生成')
            df_find = must & (df['密接/次密发现途径'].isnull())
            df[df_find].to_excel(rf'./output/质控名单（每次提醒维护后请删除里面的文件）/5.质控名单.{sick_type}发现途径维护为空.xlsx', index=False)
            print(f'{sick_type}发现途径为空名单已生成')
            sick_type = '次密'
        # 生成1+1表
        print('正在生成1+1表')
        Original().report_11()
        # time.sleep(30)  # 可delete
        print('已全部生成完毕~')

    # 7.质控
    def report_quality(self):
        df = pd.read_excel(f'./{Original().open_file()}')

    # 8.清空及创建对应文件夹
    def report_file(self):
        folder_path = './output'
        oneaddone = '/1+1表'
        important_roster = '/重要名单'
        secondary_roster = '/次要名单'
        quality_roster = '/质控名单（每次提醒维护后请删除里面的文件）'
        important_cut_roster = '/核减专属名单'
        try:
            shutil.rmtree(folder_path)
        except:
            print('清空失败,可能你本来就清空了或本来就没有这个文件夹~')
        if not os.path.exists(folder_path):
            os.makedirs(folder_path)  # 创建文件夹
            os.makedirs(folder_path + oneaddone)
            os.makedirs(folder_path + important_roster)
            os.makedirs(folder_path + secondary_roster)
            os.makedirs(folder_path + important_roster + important_cut_roster)
            os.makedirs(folder_path + quality_roster)



if __name__ == '__main__':
    Test = Original()
    print('正在清理上版数据的文件')
    Test.report_file()
    print('清空完毕')
    print(f'共有{len(Test.report)}版数据')
    # 批量汇报(使用批量汇报xlsx)
    for i in range(0, len(Test.report)):
        if Test.report['数据处理'].values[i] == '是':
            Test.dealwith()
        # 文字报数详细版
        if Test.report['详细版'].values[i] == '是':
            name = Test.report['name'].values[i]
            title = Test.report['title'].values[i]
            write = Test.report['打印名单'].values[i]
            lastnameif = Test.report['是否和lastname对比新增人数'].values[i]
            name_last = Test.report['lastname'].values[i]
            Test.report_pronusual()  # 详细版
        # 文字报数简版
        if Test.report['简版'].values[i] == '是':
            name = Test.report['name'].values[i]
            title = Test.report['title'].values[i]
            Test.report_simpleusual()
        # 新1+1表
        if Test.report['1+1表'].values[i] == '是':
            name = Test.report['name'].values[i]
            title = Test.report['title'].values[i]
            Test.report_new11()
        # 新重点场所汇报
        if Test.report['重点场所'].values[i] == '是':
            name = Test.report['name'].values[i]
            title = Test.report['title'].values[i]
            Test.report_place()
    for paragraph in Original().doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "微软雅黑"
            r = run._element.rPr.rFonts
            r.set(qn('w:eastAsia'), '微软雅黑')
    Original().doc.save(f'./output/文字报数【{Original().open_file()[6:-5]}】.docx')
    print('已完全汇报完毕!~')

